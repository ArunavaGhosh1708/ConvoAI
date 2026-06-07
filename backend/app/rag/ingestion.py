"""Document ingestion pipeline: load → chunk → embed → store."""

import logging
import uuid
from pathlib import Path
from typing import BinaryIO

from openai import AsyncOpenAI
from sqlalchemy import update
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.knowledge import Document, KnowledgeChunk

logger = logging.getLogger(__name__)

_openai = AsyncOpenAI(api_key=settings.openai_api_key)


# ---------------------------------------------------------------------------
# Loading
# ---------------------------------------------------------------------------

def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _load_docx(path: Path) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _load_html(path: Path) -> str:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(path.read_bytes(), "lxml")
    return soup.get_text(separator="\n", strip=True)


_LOADERS = {
    "pdf": _load_pdf,
    "docx": _load_docx,
    "html": _load_html,
    "txt": _load_text,
}


def load_document(path: Path, file_type: str) -> str:
    loader = _LOADERS.get(file_type)
    if loader is None:
        raise ValueError(f"Unsupported file type: {file_type}")
    return loader(path)


# ---------------------------------------------------------------------------
# Chunking  (recursive character splitter)
# ---------------------------------------------------------------------------

def _split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """Recursive character text splitter matching LangChain's behaviour."""
    separators = ["\n\n", "\n", ". ", " ", ""]
    return _recursive_split(text, separators, chunk_size, chunk_overlap)


def _recursive_split(
    text: str,
    separators: list[str],
    chunk_size: int,
    chunk_overlap: int,
) -> list[str]:
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    sep = separators[0]
    remaining_seps = separators[1:]

    if sep == "":
        # character-level fallback — overlap is baked into the stride
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start += chunk_size - chunk_overlap
        return [c for c in chunks if c.strip()]

    parts = text.split(sep)

    # Separator not present — delegate entirely to the next separator without
    # applying overlap at this level (prevents multi-level overlap accumulation).
    if len(parts) == 1:
        if remaining_seps:
            return _recursive_split(text, remaining_seps, chunk_size, chunk_overlap)
        return [text]

    chunks = []
    current = ""

    for part in parts:
        candidate = (current + sep + part) if current else part
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)
            if len(part) > chunk_size and remaining_seps:
                chunks.extend(_recursive_split(part, remaining_seps, chunk_size, chunk_overlap))
            else:
                current = part

    if current:
        chunks.append(current)

    # apply overlap by prepending the tail of the previous chunk
    if chunk_overlap > 0 and len(chunks) > 1:
        overlapped: list[str] = [chunks[0]]
        for i in range(1, len(chunks)):
            tail = overlapped[-1][-chunk_overlap:]
            overlapped.append(tail + chunks[i])
        return [c for c in overlapped if c.strip()]

    return [c for c in chunks if c.strip()]


def chunk_text(text: str) -> list[str]:
    return _split_text(text, settings.chunk_size, settings.chunk_overlap)


# ---------------------------------------------------------------------------
# Embedding
# ---------------------------------------------------------------------------

async def embed_texts(texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts using OpenAI text-embedding-3-small."""
    if not texts:
        return []
    # OpenAI supports up to 2048 inputs per request; batch in groups of 256
    batch_size = 256
    all_embeddings: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        response = await _openai.embeddings.create(
            model=settings.embedding_model,
            input=batch,
            dimensions=settings.embedding_dimensions,
        )
        all_embeddings.extend([item.embedding for item in response.data])
    return all_embeddings


# ---------------------------------------------------------------------------
# Full ingestion pipeline
# ---------------------------------------------------------------------------

async def ingest_document(
    db: AsyncSession,
    file: BinaryIO,
    filename: str,
    file_type: str,
    extra_metadata: dict | None = None,
) -> Document:
    """
    End-to-end ingestion: persist document record, load text, chunk,
    embed, and store all knowledge_chunks. Returns the updated Document.
    """
    doc = Document(filename=filename, file_type=file_type, status="processing")
    db.add(doc)
    await db.flush()  # get doc.id without committing

    try:
        # Write upload to a temp path so loaders can open it by path
        import os
        import tempfile
        suffix = f".{file_type}"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file.read())
            tmp_path = Path(tmp.name)

        try:
            raw_text = load_document(tmp_path, file_type)
        finally:
            os.unlink(tmp_path)

        chunks_text = chunk_text(raw_text)
        if not chunks_text:
            raise ValueError("Document produced no text chunks after loading")

        logger.info("Document %s: %d chunks to embed", doc.id, len(chunks_text))
        embeddings = await embed_texts(chunks_text)

        base_metadata = extra_metadata or {}
        base_metadata["filename"] = filename

        chunks = [
            KnowledgeChunk(
                id=uuid.uuid4(),
                document_id=doc.id,
                content=text,
                embedding=emb,
                metadata_={**base_metadata, "chunk_index": idx},
            )
            for idx, (text, emb) in enumerate(zip(chunks_text, embeddings))
        ]
        db.add_all(chunks)

        await db.execute(
            update(Document)
            .where(Document.id == doc.id)
            .values(status="indexed", chunk_count=len(chunks))
        )
        await db.commit()
        await db.refresh(doc)
        logger.info("Document %s indexed with %d chunks", doc.id, len(chunks))

    except Exception as exc:
        await db.rollback()
        await db.execute(
            update(Document).where(Document.id == doc.id).values(status="failed")
        )
        await db.commit()
        logger.exception("Ingestion failed for document %s: %s", doc.id, exc)
        raise

    return doc
